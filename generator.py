import json
import os
import re
from docx import Document

# --- KONFIGURASI CUAN & SEO ---
DOMAIN_UTAMA = "spacenet.my.id"
ADSENSE_ID = "ca-pub-8957427036950408" 
# GANTI ID SLOT DI BAWAH INI DENGAN ID SLOT "DISPLAY" LU (Yang buat Header)
SLOT_ID_HEADER = "8904659749" 
ADS_TXT_CONTENT = "google.com, pub-8957427036950408, DIRECT, f08c47fec0942fa0"

# --- TEMPLATE HEADER (OPTIMIZED V10) ---
# Update: Preconnect CDN, Min-Height buat Iklan, & Meta SEO Lengkap
TEMPLATE_HEADER = """<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="description" content="{DESC}">
    <title>{TITLE} | BankSoal.id</title>
    
    <link rel="preconnect" href="https://cdn.tailwindcss.com">
    <link rel="preconnect" href="https://pagead2.googlesyndication.com">
    <link rel="preconnect" href="https://cdn.ampproject.org">
    <link rel="preconnect" href="https://cdnjs.cloudflare.com">

    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css" media="print" onload="this.media='all'">
    
    <script async custom-element="amp-auto-ads" src="https://cdn.ampproject.org/v0/amp-auto-ads-0.1.js"></script>
    <script async custom-element="amp-ad" src="https://cdn.ampproject.org/v0/amp-ad-0.1.js"></script>
</head>
<body class="bg-gray-50 text-gray-800 font-sans flex flex-col min-h-screen">

<amp-auto-ads type="adsense" data-ad-client="{ADS_ID}"></amp-auto-ads>

<nav class="bg-white border-b sticky top-0 z-50 shadow-sm">
    <div class="max-w-4xl mx-auto px-4 h-16 flex items-center justify-between">
        <a href="index.html" class="font-bold text-xl text-blue-600 flex items-center gap-2">
            <i class="fa-solid fa-book-open"></i> <span class="hidden sm:inline">BankSoal.id</span>
        </a>
        <div class="flex gap-1 md:gap-4 text-xs md:text-sm font-semibold text-gray-600">
            <a href="sd.html" class="hover:text-blue-600 px-2 py-1 rounded hover:bg-blue-50 transition">SD</a>
            <a href="smp.html" class="hover:text-blue-600 px-2 py-1 rounded hover:bg-blue-50 transition">SMP</a>
            <a href="sma.html" class="hover:text-blue-600 px-2 py-1 rounded hover:bg-blue-50 transition">SMA</a>
            <a href="smk.html" class="hover:text-blue-600 px-2 py-1 rounded hover:bg-blue-50 transition">SMK</a>
        </div>
    </div>
</nav>

<main class="max-w-4xl mx-auto px-4 py-8 flex-grow">
    
    <div class="w-full mb-8 text-center bg-white p-1 rounded border border-gray-100 overflow-hidden min-h-[280px]">
        <amp-ad width="100vw" height="320"
             type="adsense"
             data-ad-client="{ADS_ID}"
             data-ad-slot="{SLOT_ID}"
             data-auto-format="rspv"
             data-full-width="">
          <div overflow=""></div>
        </amp-ad>
    </div>
"""

TEMPLATE_SEARCH = """
    <div class="mb-8 relative">
        <div class="absolute inset-y-0 left-0 flex items-center pl-3 pointer-events-none text-gray-400">
            <i class="fa-solid fa-magnifying-glass"></i>
        </div>
        <input type="text" id="searchInput" onkeyup="cariSoal()" placeholder="Cari materi, mapel, atau kelas..." 
            class="w-full p-4 pl-10 text-sm text-gray-900 border border-gray-300 rounded-lg bg-white focus:ring-blue-500 focus:border-blue-500 shadow-sm" autocomplete="off">
    </div>

    <script>
    function cariSoal() {
        var input, filter, container, items, title, i, txtValue;
        input = document.getElementById("searchInput");
        filter = input.value.toUpperCase();
        container = document.getElementById("listSoal");
        items = container.getElementsByClassName("soal-item");

        for (i = 0; i < items.length; i++) {
            title = items[i].getElementsByTagName("h3")[0];
            desc = items[i].getElementsByTagName("p")[0];
            txtValue = title.textContent + " " + desc.textContent;
            if (txtValue.toUpperCase().indexOf(filter) > -1) {
                items[i].style.display = "";
            } else {
                items[i].style.display = "none";
            }
        }
    }
    </script>
"""

TEMPLATE_FOOTER = """
</main>
<footer class="text-center py-6 text-gray-400 text-sm border-t bg-white mt-auto">
    &copy; 2026 BankSoal.id Engine<br>
    <a href="https://github.com/DHEWAYY/bank_soal" class="text-blue-500 hover:underline text-xs">Open Source Project</a>
</footer>
</body>
</html>
"""

TEMPLATE_SOAL_ITEM = """<article class="bg-white p-6 rounded-xl shadow-sm border mb-6"><div class="flex gap-3"><span class="bg-blue-100 text-blue-700 font-bold px-3 py-1 rounded h-fit text-sm">{NO}.</span><div class="w-full"><p class="text-lg font-medium mb-4">{PERTANYAAN}</p><div class="grid grid-cols-1 md:grid-cols-2 gap-3 mb-4"><div class="p-2 border rounded hover:bg-gray-50">A. {OPSI_A}</div><div class="p-2 border rounded hover:bg-gray-50">B. {OPSI_B}</div><div class="p-2 border rounded hover:bg-gray-50">C. {OPSI_C}</div><div class="p-2 border rounded hover:bg-gray-50">D. {OPSI_D}</div></div><details><summary class="cursor-pointer text-blue-600 font-semibold text-sm">Lihat Pembahasan</summary><div class="mt-2 p-3 bg-gray-50 rounded text-sm text-gray-700"><b>Jawaban: {JAWABAN}</b><br>{PEMBAHASAN}</div></details></div></div></article>"""

# --- FUNGSI MINIFY HTML (BIAR RINGAN) ---
def minify_html(content):
    # Hapus spasi berlebih di antara tag HTML
    return re.sub(r'>\s+<', '><', content.strip())

def create_docx(data, filename_base):
    try:
        doc = Document()
        meta = data.get('meta', {})
        doc.add_heading(meta.get('judul_bab', 'Latihan Soal'), 0)
        doc.add_paragraph(f"Jenjang: {meta.get('jenjang')} | Mapel: {meta.get('mapel')} | Kelas: {meta.get('kelas')}")
        for q in data.get('soal_pg', []):
            doc.add_paragraph(f"{q['no']}. {q['tanya']}")
            doc.add_paragraph(f"A. {q['opsi_a']}  B. {q['opsi_b']}  C. {q['opsi_c']}  D. {q['opsi_d']}")
        path = f"docs/downloads/{filename_base}.docx"
        doc.save(path)
        return path.replace('docs/', '')
    except: return "#"

def write_page(filename, content):
    # Tulis file dengan Minifikasi
    with open(f'docs/{filename}', 'w', encoding='utf-8') as f:
        f.write(minify_html(content))
    print(f"✅ Halaman dibuat (Minified): docs/{filename}")

def generate():
    if not os.path.exists('docs'): os.makedirs('docs')
    if not os.path.exists('docs/downloads'): os.makedirs('docs/downloads')
    
    with open('docs/CNAME', 'w') as f: f.write(DOMAIN_UTAMA)
    with open('docs/ads.txt', 'w') as f: f.write(ADS_TXT_CONTENT)
    with open('docs/.nojekyll', 'w') as f: f.write("")

    files = [f for f in os.listdir('data') if f.endswith('.json')]
    all_materi = []

    print("🚀 Memproses Materi & Optimasi Speed...")
    for filename in files:
        try:
            with open(f'data/{filename}', 'r', encoding='utf-8') as f: data = json.load(f, strict=False)
            meta = data.get('meta', {})
            judul_bab = meta.get('judul_bab', 'Bank Soal')
            jenjang = meta.get('jenjang', 'UMUM').upper()
            
            # Deskripsi SEO Dinamis
            desc_seo = f"Latihan soal {meta.get('mapel')} Kelas {meta.get('kelas')} materi {judul_bab}. Soal pilihan ganda lengkap dengan kunci jawaban dan pembahasan."

            soal_html = ""
            for q in data.get('soal_pg', []):
                soal_html += TEMPLATE_SOAL_ITEM.format(
                    NO=q['no'], PERTANYAAN=q['tanya'], 
                    OPSI_A=q['opsi_a'], OPSI_B=q['opsi_b'], OPSI_C=q['opsi_c'], OPSI_D=q['opsi_d'],
                    JAWABAN=q['jawaban'], PEMBAHASAN=q['pembahasan']
                )
            
            link_docx = create_docx(data, filename.replace('.json', ''))
            
            full_html = TEMPLATE_HEADER.format(TITLE=judul_bab, ADS_ID=ADSENSE_ID, SLOT_ID=SLOT_ID_HEADER, DESC=desc_seo) + \
                        f"<div class='mb-6'><span class='text-xs font-bold bg-blue-100 text-blue-600 px-2 py-1 rounded'>{jenjang}</span><h1 class='text-2xl font-bold mt-2'>{judul_bab}</h1><p class='text-gray-500'>{meta.get('mapel')} - {meta.get('kelas')}</p><a href='{link_docx}' class='inline-block mt-3 text-sm text-white bg-green-600 px-4 py-2 rounded hover:bg-green-700'><i class='fa-solid fa-download mr-2'></i>Download .DOCX</a></div>" + \
                        soal_html + TEMPLATE_FOOTER
            
            out_name = filename.replace('.json', '.html')
            write_page(out_name, full_html)
            
            all_materi.append({'judul': judul_bab, 'mapel': meta.get('mapel'), 'kelas': meta.get('kelas'), 'jenjang': jenjang, 'link': out_name})

        except Exception as e: print(f"❌ Skip {filename}: {e}")

    # Generate Index Pages
    def make_index(fname, title, items, manual_desc=None):
        desc = manual_desc if manual_desc else f"Kumpulan bank soal {title} terlengkap dan gratis."
        header = TEMPLATE_HEADER.format(TITLE=title, ADS_ID=ADSENSE_ID, SLOT_ID=SLOT_ID_HEADER, DESC=desc)
        
        # BANNER LCP OPTIMIZATION (Jalur VIP untuk Gambar)
        # Kita kasih fetchpriority='high' dan dimensi eksplisit biar ga goyang (CLS)
        banner_html = ""
        if fname == "index.html":
             banner_html = """
             <div class="mb-8 rounded-xl overflow-hidden shadow-sm border border-gray-100 bg-white">
                <img src="assets/banner.png" alt="Bank Soal Digital" 
                     width="800" height="300" 
                     class="w-full h-auto object-cover" 
                     fetchpriority="high" decoding="sync">
             </div>
             """

        list_html = ""
        for m in items:
            color = "bg-gray-100 text-gray-600"
            if "SD" in m['jenjang']: color = "bg-red-100 text-red-600"
            elif "SMP" in m['jenjang']: color = "bg-blue-100 text-blue-600"
            elif "SMA" in m['jenjang']: color = "bg-gray-200 text-gray-700"
            elif "SMK" in m['jenjang']: color = "bg-orange-100 text-orange-600"
            
            list_html += f"""<a href="{m['link']}" class="soal-item block p-5 bg-white border rounded-xl hover:shadow-lg transition hover:border-blue-400 group relative"><span class="absolute top-4 right-4 text-xs font-bold px-2 py-1 rounded {color}">{m['jenjang']}</span><h3 class="font-bold text-gray-800 group-hover:text-blue-600 text-lg pr-10">{m['judul']}</h3><p class="text-sm text-gray-500 mt-2"><i class="fa-solid fa-tag mr-1"></i> {m['mapel']} - {m['kelas']}</p></a>"""

        if not list_html: list_html = "<div class='col-span-2 text-center text-gray-400 py-10'>Belum ada materi.</div>"

        content = header + banner_html + \
                  f"<div class='text-center mb-8'><h1 class='text-3xl font-bold text-gray-900 mb-2'>{title}</h1><p class='text-gray-500'>{desc}</p></div>" + \
                  TEMPLATE_SEARCH + \
                  f"<div id='listSoal' class='grid grid-cols-1 md:grid-cols-2 gap-4'>{list_html}</div>" + TEMPLATE_FOOTER
        
        write_page(fname, content)

    make_index('index.html', 'Bank Soal Digital', all_materi, "Platform latihan soal gratis SD, SMP, SMA, SMK terlengkap tanpa daftar.")
    make_index('sd.html', 'Bank Soal SD', [m for m in all_materi if 'SD' in m['jenjang']])
    make_index('smp.html', 'Bank Soal SMP', [m for m in all_materi if 'SMP' in m['jenjang']])
    make_index('sma.html', 'Bank Soal SMA', [m for m in all_materi if 'SMA' in m['jenjang']])
    make_index('smk.html', 'Bank Soal SMK', [m for m in all_materi if 'SMK' in m['jenjang']])

    print("🏁 SELESAI! Web lu sekarang udah diet (Minified) & lari kenceng (Optimized)!")

if __name__ == "__main__":
    generate()
